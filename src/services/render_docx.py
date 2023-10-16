import requests
from docx import Document
from lxml import etree
from openpyxl.reader.excel import load_workbook
from bs4 import BeautifulSoup

from src import config
from src.services.utils import king_logging


class RatstishkaAPI:
    def __init__(self, auth_token: str):
        self.headers = {
            'Host': 'www.ratsit.se',
            'Connection': 'keep-alive',
            'Accept': 'application/json, text/plain, */*',
            'Content-Type': 'application/json',
            'Cookie': f'.AspNetCore.Cookies={auth_token}'
        }

        self.URL = 'https://www.ratsit.se'

        workbook = load_workbook(config.xlsx_file)
        self.sheet = workbook.active

    def find_address(self, people_id: str) -> dict | bool:
        data = {
            'who': people_id,
            'age': ['16', '120'],
            'phoneticSearch': True,
            'page': 1,
        }

        r = requests.post(self.URL + '/api/search/combined', headers=self.headers, json=data)

        if len(r.json().get('person').get('hits')) > 1:
            return {'status_code': 500, 'msg': '⚠️Ссесия умерла, требуеться замена токена!'}

        url = r.json().get('person').get('hits')[0].get('personUrl')
        r = requests.get(url, headers=self.headers)
        page = BeautifulSoup(r.content, 'html.parser')
        dom = etree.HTML(str(page))

        try:
            elements = page.find('div', class_='col-12 col-lg-6 mt-2 mt-md-0 position-relative').find_all('p')
            address = ''
            post_code = ''
            city = ''

            for element in elements:
                if element.find('span') is not None and element.find('span').text.strip() == 'Gatuadress:':
                    address = element.text.replace('Gatuadress: ', '').strip()

                if element.find('span') is not None and element.find('span').text.strip() == 'Postnummer:':
                    post_code = element.text.replace('Postnummer: ', '').strip()

                if element.find('span') is not None and element.find('span').text.strip() == 'Postort:':
                    city = element.text.replace('Postort: ', '').strip()

            full_name = dom.xpath('//*[@id="gtm-personrapport"]/div/div[2]/div[3]/div[4]/div[1]/p[6]/text()')[0].strip()
            if full_name == '':
                full_name = dom.xpath(
                    '//*[@id="gtm-personrapport"]/div/div[2]/div[3]/div[4]/div[1]/p[7]/text()'
                )[0].strip()

            if None not in [address, post_code, city, full_name]:
                print()
                king_logging(url)

                return {'address': address, 'city': post_code + ' ' + city, 'full_name': full_name}

            else:
                return {'status_code': 400}

        except IndexError:
            return {'status_code': 400}

    def auto_fill_docx(self):
        counter = 1

        while True:
            try:
                people_id = self.sheet[f'A{counter}'].value

                if people_id is not None:
                    people_info = self.find_address(people_id)

                    if people_info.get('status_code') == 400:
                        counter += 1
                        continue

                    if people_info.get('status_code') == 500:
                        return king_logging(people_info.get('msg'), type_='error')

                    address = people_info.get('address')
                    city = people_info.get('city')
                    full_name = people_info.get('full_name')

                    template_letter = Document(config.template_letter)
                    template_envelope = Document(config.template_envelope)

                    for paragraph in template_letter.paragraphs:
                        paragraph.text = paragraph.text.replace('{{ full_name }}', full_name)

                    template_letter.save(f'{config.saves_folder}{full_name}_letter.docx')
                    king_logging(f'Сделал: {full_name}_letter.docx')

                    for table in template_envelope.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for paragraph in cell.paragraphs:
                                    if '{{ full_name }}' in paragraph.text:
                                        paragraph.text = paragraph.text.replace('{{ full_name }}', full_name)
                                    if '{{ address }}' in paragraph.text:
                                        paragraph.text = paragraph.text.replace('{{ address }}', address)
                                    if '{{ city }}' in paragraph.text:
                                        paragraph.text = paragraph.text.replace('{{ city }}', city)

                    template_envelope.save(f'{config.saves_folder2}{full_name}_envelope.docx')
                    king_logging(f'Сделал: {full_name}_envelope.docx')

                counter += 1

            except ValueError:
                break

    def start(self):
        king_logging('Программа запущена.')
        self.auto_fill_docx()
        king_logging('Программа завершила работу.')
